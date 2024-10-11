# -*- coding: utf-8 -*--
from pyramid.view import view_config
import pyramid.httpexceptions as exc
from sqlalchemy import and_, text, desc
from sqlalchemy.sql.expression import func

from infolica.exceptions.custom_error import CustomError
from infolica.models.models import Affaire, Cadastre, GeosBalance, Numero, NumeroEtatHisto, AffaireNumero
from infolica.scripts.utils import Utils
from infolica.scripts.mailer import send_mail
from infolica.scripts.authentication import check_connected
import os
import json
import re
from docx.api import Document
from datetime import datetime


@view_config(route_name="balance_generate_table", request_method="GET", renderer="json")
def balance_generate_table_view(request):
    """
    Generate table balance
    """
    # Check connected
    if not check_connected(request):
        raise exc.HTTPForbidden()

    send_mail(request, [request.registry.settings["infolica_cron_mail_trigger"]], "Génération Balance Infolica\nMail généré automatiquement", "Infolica-Balance")
    return "ok"


@view_config(route_name="balance_mutation_names", request_method="GET", renderer="json")
def balance_mutation_names_view(request):
    """
    Get balance by mutation names
    """
    # Check connected
    if not check_connected(request):
        raise exc.HTTPForbidden()

    query = request.dbsession.query(GeosBalance.mutation).group_by(GeosBalance.mutation).all()

    query = [i[0] for i in query]
    return json.dumps(query)


@view_config(route_name="balance_by_affaire_id", request_method="GET", renderer="json")
def balance_view(request):
    """
    Return balance ef affaire from table GEOS
    """
    # Check connected
    if not check_connected(request):
        raise exc.HTTPForbidden()

    mutation_name = request.params["mutation_name"] if "mutation_name" in request.params else None

    if not mutation_name:
        raise CustomError(CustomError.RECORD_WITH_ID_NOT_FOUND.format(GeosBalance.__tablename__, mutation_name))

    query = request.dbsession.query(GeosBalance).filter(GeosBalance.mutation == mutation_name).all()

    return Utils.serialize_many(query)


@view_config(route_name="balance_check_existing_oldBF", request_method="POST", renderer="json")
def balance_check_existing_oldBF_new_view(request):
    """
    Check if oldBF already exist in DB and create it otherwise
    """
    # Check connected
    if not Utils.has_permission(request, request.registry.settings["affaire_numero_edition"]):
        raise exc.HTTPForbidden()

    affaire_id = request.params["affaire_id"]
    oldBF = json.loads(request.params["oldBF"])

    # Control existance of each oldBF
    numero_obj = []
    for bf in oldBF:
        bf_cadastre_id, bf_numero = bf.split("_")
        numero = request.dbsession.query(Numero).filter(and_(Numero.cadastre_id == bf_cadastre_id, Numero.numero == bf_numero)).first()

        # Create number if it doesn't exist
        if not numero:
            numero = Numero(cadastre_id=bf_cadastre_id, type_id=request.registry.settings["numero_bf_id"], numero=bf_numero, etat_id=request.registry.settings["numero_vigueur_id"])

            request.dbsession.flush()

        # Add numero to array of Numeros created
        numero_obj.append(numero)

        # Add numero_affaire link
        affnum_exists = request.dbsession.query(AffaireNumero).filter(and_(AffaireNumero.numero_id == numero.id, AffaireNumero.affaire_id == affaire_id)).first()

        if affnum_exists is None:
            affNum = AffaireNumero(affaire_id=affaire_id, numero_id=numero.id, type_id=request.registry.settings["numero_bf_id"], actif=True)
            request.dbsession.add(affNum)

        # Add numero_etat_histo link
        numEtatHisto_exists = request.dbsession.query(NumeroEtatHisto).filter(and_(NumeroEtatHisto.numero_id == numero.id, NumeroEtatHisto.numero_etat_id == request.registry.settings["numero_vigueur_id"])).first()

        if numEtatHisto_exists is None:
            numEtatHisto = NumeroEtatHisto(numero_id=numero.id, numero_etat_id=request.registry.settings["numero_vigueur_id"], date=datetime.now().date())
            request.dbsession.add(numEtatHisto)

    return Utils.serialize_many(numero_obj)


@view_config(route_name="balance_files_by_affaire_id", request_method="GET", renderer="json")
def get_balance_files_view(request):
    """
    Return balance
    """
    # Check connected
    if not check_connected(request):
        raise exc.HTTPForbidden()

    affaire_id = request.params["affaire_id"] if "affaire_id" in request.params else None
    affaires_directory = request.registry.settings["affaires_directory"]
    balance_file_rel_path = request.registry.settings["balance_file_rel_path"]
    balance_filename_prefix = request.registry.settings["balance_filename_prefix"]

    # Get affaire path and search file
    query = request.dbsession.query(Affaire).filter(Affaire.id == affaire_id).first()
    relpath = os.path.normcase(os.path.join(query.chemin, balance_file_rel_path))
    path = os.path.normcase(os.path.join(affaires_directory, relpath))

    files = []
    for filename in os.listdir(path):
        if filename.startswith(balance_filename_prefix) and (filename.endswith(".doc") or filename.endswith(".docx")):
            files.append({"filename": filename, "filepath": os.path.join(relpath, filename)})

    return files


def _format_numero_balance(cadastre_id, number):
    result = None
    if number.isnumeric():
        result = f"{cadastre_id}_{number}"
    elif "dp" in number.lower():
        result = "DP"
    return result


def _find_keyword_in_header(file_path, keyword):
    # open word document
    doc = Document(file_path)

    # go through sections to access headers
    for section in doc.sections:
        header = section.header  # get header

        # go through paragraphs
        for para in header.paragraphs:
            text = para.text.lower()
            # search keyword
            if keyword.lower() in text:
                # extract cadastre
                test = re.match("(cadastre ((de)|(du) )?)", text)
                cadastre_ = text[test.span()[1] :].strip()

                return cadastre_

    raise CustomError(CustomError.CADASTRE_NOT_FOUND_IN_FILE.format(input_file))


@view_config(route_name="balance_from_file", request_method="POST", renderer="json")
def balance_from_file_view(request):
    """
    Return balance
    """
    # Check connected
    if not check_connected(request):
        raise exc.HTTPForbidden()

    input_file = request.params["filepath"] if "filepath" in request.params else None
    affaire_id = request.params["affaire_id"] if "affaire_id" in request.params else None

    # Build complete filepath
    input_filepath = os.path.normcase(os.path.join(request.registry.settings["affaires_directory"], input_file))


    # get cadastre from balance file if exists
    try:
        cadastre_ = _find_keyword_in_header(input_filepath, "cadastre")

        result = request.dbsession.query(
            Cadastre,
            func.similarity(Cadastre.nom, cadastre_).label("similarity")
        ).order_by(desc("similarity")).limit(1).first()

        cadastre = result[0]

    except:
        # get cadastre from affaire
        cadastre = request.dbsession.query(Cadastre).join(Affaire, Affaire.cadastre_id == Cadastre.id).filter(Affaire.id == affaire_id).first()

    # get balance table
    doc = Document(input_filepath)

    table = None
    for table in doc.tables:
        if "ancien" in table.rows[0].cells[0].text:
            break

    lastBF = []
    balance = []
    for i, row in enumerate(table.rows[1:]):

        text = list(cell.text for cell in row.cells)

        if i == 0:
            lastBF = text
            continue

        for j, text_i in enumerate(text):
            if j >= 2 and text_i.isnumeric() and not text[0] == "":
                old_bf = _format_numero_balance(cadastre.id, text[0])
                new_bf = _format_numero_balance(cadastre.id, lastBF[j])
                balance.append([old_bf, new_bf])

    return balance


# Route used by external applications
@view_config(route_name="balance_from_affaire_id", request_method="GET", renderer="json")
def get_balance_from_affaire_id(request):
    """
    Send entire balance of affaire_id in a 2d array format
    """

    affaire_id = request.params["division_id"] if "division_id" in request.params else None
    cadastre_id = int(request.params["cadastre_id"]) if "cadastre_id" in request.params else None

    if affaire_id is None:
        return {"error": {"status_code": 400, "detail": "L'identifiant de l'affaire manque ! (affaire_id = {})".format(affaire_id)}}

    # get affaire
    affaire = request.dbsession.query(Affaire)
    if affaire_id.isnumeric():
        affaire = affaire.filter(Affaire.id == affaire_id)
    else:
        affaire_no_access = ""
        affaire_split = affaire_id.split("_")
        if len(affaire_split) == 1:
            last_char = "A"
            for char in affaire_id:
                char = int(char) if char.isdigit() else char
                if type(last_char) == type(char):
                    affaire_no_access += str(char)
                else:
                    affaire_no_access += "_" + str(char)
                last_char = char

            affaire_no_access += "_0"
        elif len(affaire_split) == 2:
            affaire_no_access = affaire_id + "_0"
        elif len(affaire_split) == 3:
            affaire_no_access = affaire_id
        else:
            raise exc.HTTPBadRequest()

        affaire = affaire.filter(Affaire.no_access.ilike(affaire_no_access))
    affaire = affaire.first()

    if affaire is None:
        return {"error": {"status_code": 400, "detail": "Cette affaire n'existe pas ! (affaire_id = {})".format(affaire_id)}}

    if affaire.type_id not in [1, 6, 16, 17, 19]:
        return {"error": {"status_code": 400, "detail": "Le type de cette affaire n'est pas une division ! (affaire_id = {})".format(affaire_id)}}

    if affaire.date_envoi is None:
        return {"error": {"status_code": 400, "detail": "Cette affaire n'a pas encore été envoyée ! (affaire_id = {})".format(affaire_id)}}

    if affaire.cadastre_id != cadastre_id:
        return {"error": {"status_code": 400, "detail": "Cette affaire est enregistrée sur un autre cadastre ! (affaire_id = {})".format(affaire_id)}}

    # get number relations
    sql = " \
        select nb.no_access as source, \
            na.no_access as destination \
        from infolica.numero_relation nr, \
            infolica.numero nb, \
            infolica.numero na \
        where nr.numero_id_base = nb.id \
            and nr.numero_id_associe = na.id \
            and nr.relation_type_id = 1 \
            and nr.affaire_id = {}; \
    ".format(
        affaire.id
    )

    results = request.dbsession.execute(text(sql))

    source_bf = []
    destination_bf = []
    relations = []

    for res in results:
        if res.destination:
            source_bf.append(res.source) if not res.source in source_bf else None
            destination_bf.append(res.destination) if not res.destination in destination_bf else None
            relations.append([res.source, res.destination])

    # rename "0_1" into "DP" and rename "0_2" into "RP"
    source_bf = ["DP" if x == "0_1" else x for x in source_bf]
    destination_bf = ["DP" if x == "0_1" else x for x in destination_bf]
    relations = [["DP" if y == "0_1" else y for y in x] for x in relations]
    source_bf = ["RP" if x == "0_2" else x for x in source_bf]
    destination_bf = ["RP" if x == "0_2" else x for x in destination_bf]
    relations = [["RP" if y == "0_2" else y for y in x] for x in relations]

    source_bf.sort()
    destination_bf.sort()

    balance = None

    if len(source_bf) > 0 and len(destination_bf) > 0:
        # initialize balance 2d list
        balance = [[" " for j in range(len(destination_bf) + 1)] for i in range(len(source_bf) + 1)]

        # 1rst line and row for headers
        for i in range(len(source_bf)):
            balance[i + 1][0] = source_bf[i]

        for j in range(len(destination_bf)):
            balance[0][j + 1] = destination_bf[j]

        for rel in relations:
            src_idx = source_bf.index(rel[0]) + 1
            dst_idx = destination_bf.index(rel[1]) + 1
            balance[src_idx][dst_idx] = "X"

    return {"balance": balance}
